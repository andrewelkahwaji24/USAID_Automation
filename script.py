from openpyxl import load_workbook, Workbook
from docx import Document
from email.message import EmailMessage
import smtplib
import docx2pdf
import os

def fill_word_template(template_path, output_path, data):
    try:
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found at '{template_path}'")

        document = Document(template_path)

        # Iterate over each dictionary in the data list
        for row_data in data:
            # Replace placeholders in the document paragraphs
            for paragraph in document.paragraphs:
                for key, value in row_data.items():
                    if isinstance(value, (int, float)):
                        value = str(value)  # Convert numeric values to string
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, str(value))

            # Replace placeholders in the tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in row_data.items():
                            if isinstance(value, (int, float)):
                                value = str(value)  # Convert numeric values to string
                            if key in cell.text:
                                cell.text = cell.text.replace(key, str(value))

        # Remove the existing file if it exists
        if os.path.exists(output_path):
            os.remove(output_path)

        # Save the filled document
        document.save(output_path)
        print("Filled document saved successfully:", output_path)
        return True
    except Exception as e:
        print(f"Error during template filling: {e}")
        return False

def convert_to_pdf(docx_path, pdf_path):
    try:
        docx2pdf.convert(docx_path, pdf_path)
        print("PDF generated successfully:", pdf_path)
        return True
    except Exception as e:
        print(f"Error during PDF conversion: {e}")
        return False

def send_email_with_attachment(sender_name, sender_email, recipient_email, subject, body, attachment_path):
    try:
        # Create EmailMessage object
        msg = EmailMessage()
        msg['Subject'] = subject
        msg['From'] = f"{sender_name} <{sender_email}>"
        msg['To'] = recipient_email
        msg.set_content(body)

        # Attach the file
        with open(attachment_path, 'rb') as f:
            file_data = f.read()
            file_name = os.path.basename(attachment_path)
        msg.add_attachment(file_data, maintype='application', subtype='octet-stream', filename=file_name)

        # Send the email
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
            smtp.login(sender_email, 'nixt lulb dojj dqwt')  # Replace with your email password
            smtp.send_message(msg)
        print("Email sent successfully.")
        return True
    except Exception as e:
        print(f"Error: {e}")
        return False

# Ensure output directory exists
output_dir = 'output'
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# Load data from Excel and calculate specific hours
try:
    wb = load_workbook('data.xlsx', data_only=True)
    sheet = wb.active

    # Dictionary to store total hours for each volunteer
    hours_summary = {}

    # Iterate over rows in the Excel sheet
    for row in sheet.iter_rows(min_row=2, values_only=True):
        name, ht, h1, h2, h3, h4, tasks, feedback, email = row  # Include email in row structure

        # Calculate total hours for h1, h2, h3, h4
        total_hours = sum(filter(None, [h1, h2, h3, h4]))  # Summing non-empty hours

        # Update hours summary
        if name in hours_summary:
            hours_summary[name] += total_hours
        else:
            hours_summary[name] = total_hours

        # Create a dictionary for the current row
        row_data = {
            '{{name}}': name,
            '{{ht}}': ht,
            '{{h1}}': h1,
            '{{h2}}': h2,
            '{{h3}}': h3,
            '{{h4}}': h4,
            '{{tasks}}': tasks,
            '{{feedback}}': feedback
        }

        # Fill Word template
        filled_document_path = os.path.join(output_dir, f'filled_document_{name}.docx')
        template_path = 'template.docx'  # Ensure this path is correct
        if fill_word_template(template_path, filled_document_path, [row_data]):
            # Convert Word document to PDF
            pdf_path = os.path.join(output_dir, f'filled_document_{name}.pdf')
            print("Converting Word document to PDF...")
            if convert_to_pdf(filled_document_path, pdf_path):
                # Send email with attachment to the email from the sheet
                sender_name = 'Andrew El Kahwaji'  # Your name
                sender_email = 'andrewelkahwaji.coo@gmail.com'  # Replace with your email address
                subject = f'USAID Document for {name} - December  2024'
                body = f"""
                Hello 
                
                Kindly find attached your volunteering paper for the month of January 2025
                
                Sincerely yours 
                President OFFICE OF InnovaThrive
                """

                if send_email_with_attachment(sender_name, sender_email, email, subject, body, pdf_path):  # Use the non-protected PDF
                    print(f"Email sent for {name} to {email}")
                else:
                    print(f"Failed to send email for {name} to {email}.")
            else:
                print(f"PDF conversion failed for {name}.")
        else:
            print(f"Template filling failed for {name}.")

    # Create summary Excel sheet
    summary_wb = Workbook()
    summary_sheet = summary_wb.active
    summary_sheet.title = 'Volunteer Summary'

    # Write headers
    summary_sheet['A1'] = 'Volunteer Name'
    summary_sheet['B1'] = 'Total Hours'

    # Write data
    row_num = 2
    for volunteer, total_hours in hours_summary.items():
        summary_sheet[f'A{row_num}'] = volunteer
        summary_sheet[f'B{row_num}'] = total_hours
        row_num += 1

    # Save summary Excel workbook
    summary_excel_path = os.path.join(output_dir, 'volunteer_summary.xlsx')
    summary_wb.save(summary_excel_path)
    print(f"Summary Excel sheet saved successfully: {summary_excel_path}")

    # Close summary workbook
    summary_wb.close()

except Exception as e:
    print(f"Error: {e}")
